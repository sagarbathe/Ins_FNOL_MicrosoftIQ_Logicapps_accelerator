# Generates docs/End_to_End_Deployment_Guide.docx
# Run: python generate_e2e_deployment_doc.py
#
# Comprehensive end-to-end architecture + deployment runbook for the Auto
# FNOL Triage accelerator's Agent Identity solution, covering every
# component (what it does, how it was deployed, prerequisites) plus a
# sequential, copy-pasteable deployment section for a new developer.

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_callout(doc, label, text, color_hex="FFF2CC"):
    """A shaded, bordered single-cell 'table' used as a callout/note box for
    important warnings or tips, to make them visually stand out from normal
    body paragraphs."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    return table


def add_toc_field(doc):
    """Inserts a real, auto-updating Word Table of Contents field (based on
    Heading 1/2 styles) instead of a manually-typed list, so the page
    numbers and structure stay correct if headings are added/removed. Word
    will prompt to update the field (or press F9) the first time the file
    is opened - this is normal Word TOC-field behavior, not a bug."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_page_number_footer(doc):
    """Adds a 'Page X of Y' footer to every section, for easier navigation
    in a printed or scrolled-through copy of this multi-section guide."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")

    def _field(instr):
        f_begin = OxmlElement("w:fldChar"); f_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText"); instr_text.set(qn("xml:space"), "preserve"); instr_text.text = instr
        f_sep = OxmlElement("w:fldChar"); f_sep.set(qn("w:fldCharType"), "separate")
        f_end = OxmlElement("w:fldChar"); f_end.set(qn("w:fldCharType"), "end")
        r = p.add_run()
        r._r.append(f_begin); r._r.append(instr_text); r._r.append(f_sep); r._r.append(f_end)

    _field("PAGE")
    p.add_run(" of ")
    _field("NUMPAGES")



doc = Document()

# ---------- Title ----------
doc.add_heading("Auto FNOL Triage Accelerator — End-to-End Architecture & Deployment Guide", level=0)
sub = doc.add_paragraph()
r = sub.add_run(
    "Agent Identity edition: Azure AI Foundry orchestrator agent reusing Fabric IQ, Foundry IQ, and Work IQ, "
    "authenticated via a dedicated Entra ID service identity with its own mailbox and Teams membership — no OBO/Bot "
    "Framework required."
)
r.italic = True
r.font.size = Pt(12)

meta = doc.add_paragraph()
meta.add_run("Tenant: ").bold = True
meta.add_run("MngEnvMCAP146722.onmicrosoft.com\n")
meta.add_run("Subscription: ").bold = True
meta.add_run("ME-MngEnvMCAP146722-sagarbathe-1 (04054f52-6b7b-47c7-b836-005253626f42)\n")
meta.add_run("Repository: ").bold = True
meta.add_run("sagarbathe/Ins_FNOL_MicrosoftIQ_Logicapps_accelerator (local only — not yet pushed to GitHub)\n")
meta.add_run("Document date: ").bold = True
meta.add_run("2026-08-31")

add_page_number_footer(doc)

doc.add_page_break()

# ---------- Table of contents (auto-updating Word field) ----------
add_heading(doc, "Contents", level=1)
note = doc.add_paragraph()
note.add_run(
    "(This is a live Word Table of Contents field. If page numbers below show as blank or out of date, "
    "right-click it and choose \"Update Field\", or press Ctrl+A then F9 to refresh the whole document.)"
).italic = True
add_toc_field(doc)
doc.add_page_break()

# ---------- 1. Overview ----------
add_heading(doc, "1. Solution Overview", level=1)
doc.add_paragraph(
    "This accelerator replaces the original Power Automate + Copilot Studio FNOL (First Notice of Loss) intake "
    "pipeline with an Azure AI Foundry Agent Service orchestrator agent. The orchestrator reuses three existing "
    "'IQ' building blocks already deployed in this tenant:"
)
add_bullets(doc, [
    "Fabric IQ — a Fabric data agent over an insurance Ontology (Policyholder, Vehicle, Adjuster, RepairShop, "
    "Policy, Claim, FraudSignal, SubrogationFlag) exposed via Fabric's MCP endpoint.",
    "Foundry IQ — an existing governed Foundry knowledge agent (auto policy wording, coverage definitions, "
    "SIU fraud red flags, subrogation methodology), reused unchanged via a Connected Agent Tool.",
    "Work IQ — Microsoft Graph search (SharePoint/OneDrive documents and Outlook mail) for fresh, ungoverned "
    "operational content (CAT bulletins, SIU routing-change emails) not present in the governed knowledge base.",
])
doc.add_paragraph(
    "A Logic App trigger ('When a new email arrives' on the FNOL intake mailbox) creates one independent Foundry "
    "thread per email, runs the orchestrator agent, and posts the result as a new Teams channel message (the root "
    "of a per-case reply chain). A second Logic App polls for human replies in Teams and routes them back to the "
    "correct case's Foundry thread, so concurrent cases never cross-contaminate context."
)
doc.add_paragraph(
    "The key architectural decision in this iteration is the Agent Identity pattern: a dedicated Entra ID user "
    "(svc-fnol-agent@MngEnvMCAP146722.onmicrosoft.com) with its own mailbox and Teams membership, plus a dedicated "
    "public-client Entra app registration, is used to acquire delegated (user) tokens for the two capabilities that "
    "have NO application-permission equivalent in Microsoft Graph/Fabric: posting new Teams channel messages and "
    "querying the Fabric ontology data agent's MCP endpoint. See Section 4 for why this was necessary."
)

# ---------- 2. Architecture diagram (described) ----------
add_heading(doc, "2. Architecture Diagram (described)", level=1)
doc.add_paragraph(
    "A text description of the component flow (insert an actual diagram image here if desired for presentations):"
)
add_code_block(doc, """
[Outlook Inbox: FNOL intake mailbox]
        │  (new email trigger, Office 365 connector)
        ▼
[Logic App #1: workflow.json]
        │  1. Compute stable CaseId (hash of internetMessageId)
        │  2. POST /threads (Foundry) -> new thread_id
        │  3. POST email content as thread message
        │  4. POST /runs -> start orchestrator agent run
        │  5. Poll /runs/{id} until completed
        │  6. GET latest message (agent's structured summary)
        ▼
[Foundry AI Project: sbazureaimodels-project01]
        │  Orchestrator agent (auto-fnol-triage-orchestrator)
        │  Tools:
        │    - fabric_iq_ontology   (OpenAPI tool, Connection auth)  ──┐
        │    - foundry_iq_knowledge_agent (Connected Agent Tool)      │
        │    - workiq_graph_search (OpenAPI tool, Connection auth)  ──┤
        ▼                                                             │
[Logic App #1 continued]                                              │
        │  7. POST /post-new-case-alert  (Work IQ webapp) ────────────┤
        │  8. POST /insert-case-thread-map (Work IQ webapp) ──────────┤
        ▼                                                             │
[Teams Channel: Contoso / General]                                    │
   New top-level message = reply-chain root for this case              │
                                                                       │
[Logic App #2: workflow-teams-reply-poller.json]                     │
        │  Polls tracked case threads every 30 seconds                │
        │  Trigger concurrency forced to runs=1 (strictly sequential) │
        │  Finds human replies (replyToId set)                        │
        │  POST /resolve-case-reply (Work IQ webapp) ──────────────────┘
        │  -> resolves CaseId + foundry_thread_id
        │  Continues the SAME Foundry thread (follow-up run)
        │  POST /post-case-reply (nests answer under case's root msg)

[Work IQ webapp: app-autofnol-workiq.azurewebsites.net]
   Flask app, Agent-Identity-authenticated, exposes:
     /query-ontology            -> Fabric IQ MCP proxy
     /search-mail /search-documents -> Graph search proxy
     /send-email                -> Graph /me/sendMail escalation proxy
     /post-new-case-alert /post-case-reply -> Teams posting (delegated)
     /insert-case-thread-map /resolve-case-reply -> CaseThreadMap SQL CRUD

[Fabric SQL Database: AutoFNOLCaseThreadMap]
   CaseThreadMap table (CaseId PK) — correlates case <-> Teams thread <-> Foundry thread
""")

# ---------- 3. Component reference ----------
add_heading(doc, "3. Component Reference", level=1)

components = [
    (
        "Agent Identity user\n(svc-fnol-agent@MngEnvMCAP146722.onmicrosoft.com)",
        "Dedicated Entra ID user acting as the solution's own 'robot user' — has its own Exchange Online mailbox, "
        "is a member of the target Teams team, and holds Contributor on the Fabric workspace (WS_AutoFNOL only) and "
        "Foundry User on the Foundry account. All delegated (user) tokens used by the webapp are "
        "acquired against this identity via device-code sign-in, cached to disk, and silently refreshed thereafter.",
        "Created once via Entra ID admin center / Graph API (prior session). Mailbox auto-provisioned by Microsoft "
        "365 licensing. Added as a member of the Teams team via Teams admin/Graph.",
        "An M365/Entra tenant with an available user license (Exchange + Teams). Ability to add the user to the "
        "target Teams team as a member.",
    ),
    (
        "Agent Identity public-client app registration\n(FNOL-AgentIdentity-GraphPublicClient)",
        "A dedicated Entra ID app registration (public client, no secret) used ONLY to acquire delegated tokens on "
        "behalf of the Agent Identity user via MSAL device-code flow, for three resources: Microsoft Graph (Mail."
        "Read, Mail.Send, Sites.Read.All, Chat.Read, ChannelMessage.Send, ChannelMessage.Read.All), Power BI/Fabric (Item.Read.All, Item.Execute.All), "
        "and Azure SQL Database (user_impersonation). A dedicated app is required because Microsoft's own first-"
        "party public clients (e.g. the Azure CLI) cannot be granted these delegated scopes (AADSTS65002).",
        "Created via 'az ad app create' + 'az ad app permission add' for each resource's delegated scope, then "
        "admin-consented via direct Graph 'oauth2PermissionGrants' POST (the 'az ad app permission admin-consent' "
        "CLI command was unreliable for multi-resource apps in testing).",
        "Global Administrator or Privileged Role Administrator rights to grant admin consent across 3 resources.",
    ),
    (
        "Fabric IQ\n(DA_AutoFNOL_Ontology data agent)",
        "A Microsoft Fabric Data Agent published over an Ontology graph (Policyholder/Vehicle/Adjuster/RepairShop/"
        "Policy/Claim/FraudSignal/SubrogationFlag) in workspace WS_AutoFNOL, exposed via Fabric's MCP endpoint. "
        "Provides grounded, structured answers to policy/claim-specific questions.",
        "Pre-existing (built and published in an earlier phase of this accelerator, unchanged this session). The "
        "webapp's /query-ontology endpoint proxies to it using the Agent Identity's delegated Fabric ('.default') "
        "token via the mcp Python package's streamable-http client.",
        "An active (non-paused) Fabric capacity; the Agent Identity must hold at least Contributor on the Fabric "
        "workspace hosting the data agent and ontology.",
    ),
    (
        "Foundry IQ\n(existing knowledge agent, asst_thLtP20E4lDRqxiaW3rtftLF)",
        "A pre-existing, governed Foundry knowledge agent (with its own vector index / grounding data) covering "
        "policy wording, coverage-part definitions, FNOL triage-tier rules, SIU fraud red flags, state regulatory "
        "requirements, and subrogation methodology. Always returns citations.",
        "Pre-existing; referenced (not recreated) by the orchestrator via a ConnectedAgentTool pointing at its "
        "agent id (read from foundry/foundry_knowledge_agent_id.txt or FOUNDRY_KNOWLEDGE_AGENT_ID env var).",
        "The knowledge agent must already exist in the same Foundry project as the orchestrator (agent-to-agent "
        "connected tools only work within one project).",
    ),
    (
        "Work IQ\n(Graph search wrapper)",
        "A thin Python wrapper (tools_workiq_graph.py) around Microsoft Graph's Search API for SharePoint/OneDrive "
        "documents and Outlook mail plus Graph /me/sendMail, exposed as three webapp endpoints (/search-documents, "
        "/search-mail, /send-email). Used for time-sensitive operational content (CAT bulletins, SIU routing-change "
        "emails) and explicit escalation emails that the user asks the agent to send; email sending is never done "
        "automatically or by guessing recipients.",
        "Code-only component; deployed as part of the Work IQ webapp (see below). No separate Azure resource.",
        "The Agent Identity must have delegated Mail.Read, Mail.Send, Sites.Read.All, Chat.Read, "
        "ChannelMessage.Send, and ChannelMessage.Read.All Graph scopes consented on its public-client app registration.",
    ),
    (
        "Work IQ webapp\n(app-autofnol-workiq, Azure App Service Linux/Python)",
        "A small Flask app that is the single integration point exposing all Agent-Identity-authenticated "
        "capabilities as plain, API-key-protected HTTPS endpoints, so the Foundry orchestrator's OpenAPI tools and "
        "the Logic Apps can call them without needing to understand MSAL/device-code/token-cache mechanics "
        "themselves. Endpoints: GET /healthz; POST /search-documents, /search-mail, /send-email, /query-ontology, "
        "/post-new-case-alert, /post-case-reply, /resolve-case-reply, /insert-case-thread-map, /list-open-cases, "
        "/list-case-replies, /mark-reply-processed. The webapp seeds the Agent Identity token cache from chunked "
        "app settings (AGENT_IDENTITY_TOKEN_CACHE_B64_0..N plus AGENT_IDENTITY_TOKEN_CACHE_B64_COUNT) because a "
        "single App Service setting is too small for the full base64 blob; agent_identity_auth.py reassembles the "
        "chunks at runtime and still falls back to the legacy single-setting name for compatibility. Teams HTML "
        "rendering also applies inline block spacing and urgency-line colorization so replies render with visible "
        "section spacing and highlighted urgency indicators in Teams.",
        "Deployed via 'az webapp deploy --type zip' (zip containing app.py + the tools_*.py modules + "
        "agent_identity_auth.py + post_to_teams.py + requirements.txt) to an existing App Service Plan (B1 Linux). "
        "App settings (API key, Fabric/Graph IDs, SQL connection string, Agent Identity client id) are configured "
        "via 'az webapp config appsettings set' — never hardcoded in source.",
        "An Azure App Service Linux Python 3.11 plan; the Agent Identity's MSAL token cache file uploaded/seeded as "
        "chunked base64 app settings (or regenerated on the box); WORKIQ_API_KEY app setting matching the value the "
        "Foundry Connection and Logic Apps use in their 'x-api-key' header. Note: there are two copies of "
        "agent_identity_auth.py (shared/ and triggers/webapp/) and they must be kept manually in sync — changing "
        "only shared/ does not update the deployed webapp code.",
    ),
    (
        "Foundry orchestrator agent\n(auto-fnol-triage-orchestrator)",
        "The central Foundry Agent Service agent that routes each question to the right tool (Fabric IQ for named "
        "policy/claim lookups, Foundry IQ for general knowledge, Work IQ for CAT bulletins/SIU routing emails), "
        "assesses fraud/subrogation red flags, and produces a structured summary for Teams posting. Uses two "
        "OpenAPI tools (calling the Work IQ webapp, authenticated via a Foundry Connection) and one Connected "
        "Agent Tool (Foundry IQ).",
        "Created/recreated via 'python foundry/create_orchestrator_agent.py', which deletes any previous version "
        "(tracked in foundry/orchestrator_agent_id.txt) before creating a fresh one from the OpenAPI specs and "
        "instructions embedded in that script.",
        "A Foundry project with a deployed chat-completion model (FOUNDRY_MODEL_DEPLOYMENT); the "
        "workiq-api-key-conn Foundry connection (see next row) must already exist.",
    ),
    (
        "Foundry connection\n(workiq-api-key-conn, Custom Keys)",
        "A Foundry project-level connection storing the Work IQ webapp's API key as a Custom Keys credential, "
        "referenced by the orchestrator's two OpenAPI tools via OpenApiConnectionAuthDetails. This keeps the API "
        "key out of the agent definition and out of source control.",
        "Created once via a direct ARM REST PUT (the azure-ai-projects SDK's ConnectionsOperations is currently "
        "read-only — no create/update method is exposed client-side as of azure-ai-projects 2.5.0).",
        "Contributor (or higher) on the Foundry account/project to create connections via ARM.",
    ),
    (
        "CaseThreadMap store\n(Fabric SQL Database)",
        "A single small table (CaseId PK, TeamId, ChannelId, RootMessageId, FoundryThreadId, timestamps) in the "
        "same Fabric workspace, correlating each case to its Teams reply-chain root and Foundry thread so that "
        "concurrent cases and human follow-up replies are never cross-attributed (see Section 5).",
        "Table created once via a Python script authenticating with the Agent Identity's SQL-scoped Entra token "
        "(SQL_COPT_SS_ACCESS_TOKEN) and issuing CREATE TABLE + a unique index. All runtime reads/writes go through "
        "the Work IQ webapp's /insert-case-thread-map and /resolve-case-reply endpoints (pyodbc + Entra token), not "
        "a Logic App SQL connector (which cannot use Entra-token auth without another OBO hop).",
        "ODBC Driver 18 for SQL Server present in the webapp's runtime (bundled via pyodbc's Linux wheel on the App "
        "Service Python image used in this deployment); the Agent Identity must have db_datareader/db_datawriter "
        "(or equivalent) on the Fabric SQL Database; the Fabric capacity must be Active (not paused) for the "
        "database to accept connections.",
    ),
    (
        "Logic App #1\n(workflow.json — email intake)",
        "Consumption Logic App triggered on new email in the FNOL intake mailbox (Office 365 Outlook connector). "
        "Creates a Foundry thread, runs the orchestrator, posts the new-case Teams alert, and persists the "
        "CaseThreadMap row — all via HTTP actions against the Foundry REST API (Managed Identity auth) and the "
        "Work IQ webapp (API-key auth).",
        "Deployed as an Azure Logic App (Consumption) resource; workflow.json is the definition, "
        "parameters.json.template / connections.json.template are filled in per-environment at deploy time.",
        "An Office 365 API connection (Outlook) authorized against the FNOL intake mailbox; the Logic App's "
        "Managed Identity must have the Foundry User role on the Foundry account (for the "
        "ManagedServiceIdentity-authenticated Foundry REST calls). Cognitive Services User alone is insufficient "
        "for the current Foundry data plane; keep both roles only if other parts of your environment still depend "
        "on the legacy Cognitive Services User grant.",
    ),
    (
        "Logic App #2\n(workflow-teams-reply-poller.json — reply handling)",
        "A second, independent Consumption Logic App on a 30-second recurrence that polls the tracked case threads "
        "for new human replies, resolves each reply's case via the Work IQ webapp's /resolve-case-reply endpoint, "
        "continues that case's existing Foundry thread with a follow-up run, and posts the agent's answer nested "
        "under the original case's root message (never as a new top-level message). Its trigger sets "
        "runtimeConfiguration.concurrency.runs = 1 so poll cycles run strictly sequentially and cannot overlap on "
        "the same Foundry thread, preventing intermittent 400 'Can't add messages to thread while a run is active' "
        "errors.",
        "Deployed the same way as Logic App #1 (separate Consumption Logic App resource, own recurrence trigger, "
        "no email trigger dependency).",
        "The Logic App's Managed Identity needs Foundry User on the Foundry account plus Graph application "
        "permission ChannelMessage.Read.All (app-only IS supported for reading, unlike sending) to poll channel "
        "message replies.",
    ),
]

for name, what, how, prereq in components:
    add_heading(doc, name, level=2)
    p = doc.add_paragraph(); p.add_run("What it does: ").bold = True; p.add_run(what)
    p = doc.add_paragraph(); p.add_run("How it was deployed: ").bold = True; p.add_run(how)
    p = doc.add_paragraph(); p.add_run("Prerequisites: ").bold = True; p.add_run(prereq)

doc.add_page_break()

# ---------- 4. Why Agent Identity ----------
add_heading(doc, "4. Why an Agent Identity (Not OBO / Bot Framework)", level=1)
doc.add_paragraph(
    "Two capabilities in this solution have NO application-permission (app-only/service-principal) equivalent in "
    "Microsoft Graph or Fabric, and therefore cannot be done with a plain Managed Identity or client-secret service "
    "principal, no matter how it is configured:"
)
add_table(doc,
    ["Capability", "Why app-only auth cannot work", "Agent Identity fix"],
    [
        ("Post a new Teams channel message",
         "Microsoft Graph has no application permission equivalent to ChannelMessage.Send — sending channel "
         "messages is delegated-only.",
         "Acquire a delegated Graph token for the Agent Identity user (device-code-bootstrapped, cached, silently "
         "refreshed) and call the Graph API as that user."),
        ("Query the Fabric ontology data agent via MCP",
         "The Fabric MCP data-agent endpoint in this tenant only accepts delegated tokens; app-only Fabric tokens "
         "were rejected in testing.",
         "Acquire a delegated Fabric ('.default') token for the same Agent Identity user."),
    ],
    col_widths=[1.8, 2.6, 2.6],
)
doc.add_paragraph(
    "An alternative considered earlier was a Bot Framework registration using the On-Behalf-Of (OBO) flow, where "
    "a Teams bot's OAuth Token Store captures a real signed-in user's consent and the backend exchanges that token "
    "for downstream resources. That approach works but requires: (a) registering and provisioning a Bot Framework "
    "channel registration, (b) a public bot messaging endpoint for OAuth callbacks, and (c) an actual human user to "
    "complete an interactive consent at least once (and periodically thereafter, depending on refresh-token "
    "lifetime policy). See docs/OBO_Bot_Teams_Design.docx for the full design and comparison."
)
doc.add_paragraph(
    "The Agent Identity approach avoids all of that: there is no bot registration, no public callback endpoint, "
    "and no dependency on a specific human's consent — the mailbox and Teams membership belong to the solution "
    "itself. The tradeoff is that the Agent Identity's own token cache must be bootstrapped once (interactively, "
    "via device code) and its access must be provisioned like any other tenant member (mailbox license, Teams "
    "membership, Fabric workspace role) — a one-time administrative setup rather than an ongoing per-user consent."
)

# ---------- 5. Concurrency design ----------
add_heading(doc, "5. Concurrency & Teams Thread Isolation Design", level=1)
doc.add_paragraph(
    "Full design detail lives in docs/teams-concurrency-design.md; summarized here for completeness. The core "
    "guarantee: N concurrent FNOL emails produce N independent Foundry threads and N independent Teams reply "
    "chains, and any human reply — no matter how many other cases are in flight — is deterministically routed back "
    "to the correct case."
)
add_bullets(doc, [
    "Every new case is posted as a brand-new TOP-LEVEL Teams channel message (never appended to an existing one) "
    "— Teams natively renders all replies to that message as a separate, visually isolated thread.",
    "A CaseId is derived deterministically from the email's internetMessageId (base64-hashed), making the whole "
    "pipeline naturally idempotent/safe to retry without creating duplicate cases or threads.",
    "A CaseThreadMap row (CaseId PK, TeamId, ChannelId, RootMessageId, FoundryThreadId) is inserted idempotently "
    "(duplicate CaseId insert is rejected by the primary key) immediately after posting the new case alert.",
    "Human replies are correlated via Teams' own replyToId field — GET .../messages/{root}/replies (or a delta "
    "poll of the whole channel) reveals which root message a reply belongs to, which is looked up directly against "
    "CaseThreadMap(TeamId, ChannelId, RootMessageId) with zero ambiguity.",
    "Follow-up runs ALWAYS reuse the case's existing foundry_thread_id (never mint a new thread for a reply), so "
    "the orchestrator agent retains full prior context for that case.",
    "The reply-poller trigger runs every 30 seconds with runtimeConfiguration.concurrency.runs = 1, so if one poll "
    "cycle is still waiting on a Foundry run, the next recurrence is queued instead of racing to post another "
    "message into the same thread.",
])
doc.add_paragraph(
    "This was verified with an automated end-to-end concurrency test: two simultaneous FNOL 'emails' (Cases A and "
    "B) were run in parallel through the full pipeline (Foundry thread creation → orchestrator run → Teams post → "
    "CaseThreadMap insert). Result: two distinct Foundry thread IDs and two distinct Teams root message IDs were "
    "produced with no collision, and each case's orchestrator response correctly reflected its own policy data "
    "(POL-00005 vs. POL-00012) with no cross-contamination. See Section 7 for the full verification checklist."
)

doc.add_page_break()

# ---------- 6. Deployment runbook ----------
add_heading(doc, "6. Sequential Deployment Runbook", level=1)
doc.add_paragraph(
    "All the steps below have been consolidated into a single PowerShell orchestration script, "
    "deploy_solution.ps1, at the repository root, so a developer does not need to copy/paste individual az CLI "
    "commands by hand. The script runs every step in order and pauses at the few points that genuinely require a "
    "human (interactive Agent Identity sign-in, portal-only actions, or filling in environment-specific values)."
)

add_heading(doc, "6.0 Required .env values — what they are and where to find them", level=2)
doc.add_paragraph(
    ".env.example is organized into two sections: Section 1 lists every value you MUST supply yourself before "
    "running the script (no working default exists), and Section 2 lists values that already have a sensible "
    "default and only need changing if you want non-default names/regions. Copy .env.example to .env (Step 0 does "
    "this for you) and fill in Section 1 using the table below."
)
add_table(
    doc,
    ["Variable", "What it is", "Where to find it"],
    [
        ("AZURE_SUBSCRIPTION_ID", "Subscription hosting your resources.",
         "az account show --query id -o tsv, or Portal → Subscriptions."),
        ("FOUNDRY_RESOURCE_GROUP / FOUNDRY_ACCOUNT_NAME / FOUNDRY_PROJECT_NAME",
         "The existing Foundry (AI Services) account + project you are reusing.",
         "Portal → Azure AI Foundry → your account → Overview (resource group + account name); project name is "
         "the project you opened inside that account (also visible in the project's Overview page)."),
        ("FOUNDRY_PROJECT_ENDPOINT", "The project's API endpoint URL.",
         "Azure AI Foundry portal → your project → Overview → 'Azure AI Foundry project endpoint', or "
         "https://<account>.cognitiveservices.azure.com/api/projects/<project>."),
        ("FOUNDRY_KNOWLEDGE_AGENT_ID (or FOUNDRY_KNOWLEDGE_AGENT_ID_FILE)",
         "The id of your EXISTING Foundry IQ knowledge agent that you are reusing (not created by this script).",
         "Azure AI Foundry portal → your project → Agents → open the Foundry IQ knowledge agent → copy its Agent "
         "ID from the details pane; or query it with the azure-ai-projects SDK's client.agents.list()."),
        ("FABRIC_WORKSPACE_ID", "The Fabric workspace hosting Fabric IQ's ontology data agent.",
         "https://app.fabric.microsoft.com → open the workspace → Workspace settings → copy the GUID from the "
         "browser URL (…/groups/<FABRIC_WORKSPACE_ID>/…)."),
        ("FABRIC_DATA_AGENT_ID", "The published Ontology data agent (Fabric IQ) you are reusing.",
         "https://app.fabric.microsoft.com → open the workspace → open the Data agent item → Settings/details "
         "pane → copy its item id from the URL (…/data-agents/<FABRIC_DATA_AGENT_ID>/…)."),
        ("CASETHREADMAP_FABRIC_SQL_ENDPOINT / CASETHREADMAP_FABRIC_SQL_DATABASE / "
         "CASETHREADMAP_FABRIC_SQL_CONNECTION_STRING",
         "The Fabric SQL Database (in the same workspace) used to store case/thread mappings.",
         "https://app.fabric.microsoft.com → open your Fabric SQL Database item → Settings (gear icon) → "
         "Connection strings → copy the Server (endpoint), Database name, and full ODBC connection string shown "
         "there."),
        ("WORKIQ_RESOURCE_GROUP / LOGICAPP_RESOURCE_GROUP",
         "Resource group(s) the script deploys the Work IQ webapp and Logic Apps into.",
         "Pick an existing resource group (Portal → Resource groups) or a new name — Step 5/Step 8 auto-create "
         "the group and resources inside it if they don't already exist."),
        ("WORKIQ_API_KEY", "A shared secret YOU invent — not looked up anywhere. It is not auto-generated by any "
         "script; you must set it yourself before Step 5 runs.",
         "Generate your own random string, e.g. run: python -c \"import secrets;print(secrets.token_urlsafe(24))\" "
         "and paste the result here. This same value is pushed as an app setting on the Work IQ webapp (Step 5), "
         "stored in a Foundry Custom Keys connection (Step 6), and referenced by the Logic Apps (Step 8) — it "
         "must match everywhere, which the script guarantees automatically since all three steps read it from "
         "this single .env value."),
        ("AAD_TENANT_ID / AAD_CLIENT_ID / AAD_CLIENT_SECRET",
         "An app registration used for Graph app-only calls (service_principal auth mode fallback).",
         "Portal → Microsoft Entra ID → App registrations → (create or reuse one) → Overview for tenant/client "
         "id; Certificates & secrets → New client secret for the secret value. Tenant id also via: "
         "az account show --query tenantId -o tsv."),
        ("AGENT_IDENTITY_TENANT_ID / AGENT_IDENTITY_UPN",
         "The tenant id and UPN of the dedicated Agent Identity user created in Step 1.",
         "Tenant id: az account show --query tenantId -o tsv. UPN: whatever you type in when Step 1 prompts "
         "'Enter the Agent Identity's UPN' (e.g. svc-fnol-agent@yourtenant.onmicrosoft.com)."),
        ("AGENT_IDENTITY_TOKEN_CACHE_PATH", "Local file path where the Agent Identity's MSAL token cache is "
         "persisted after Step 3's device-code sign-in.",
         "Any writable path on your machine, e.g. .secrets\\agent_identity_token_cache.bin under the repo root "
         "(the .secrets folder is already gitignored)."),
        ("AGENT_IDENTITY_PUBLIC_CLIENT_ID", "The public-client app registration used for the Agent Identity's "
         "device-code sign-in.",
         "Printed by Step 2 after it creates 'FNOL-AgentIdentity-GraphPublicClient', or look it up at Portal → "
         "Microsoft Entra ID → App registrations → FNOL-AgentIdentity-GraphPublicClient → Overview → Application "
         "(client) ID."),
        ("AGENT_IDENTITY_OBJECT_ID", "The Entra object id of the Agent Identity user, used by the reply poller to "
         "ignore the agent's own Teams posts.",
         "Run: az ad user show --id <AGENT_IDENTITY_UPN> --query id -o tsv, or Portal → Microsoft Entra ID → Users "
         "→ open the Agent Identity user → Object ID."),
        ("TEAMS_TEAM_ID / TEAMS_CHANNEL_ID", "The Teams team + channel where case alerts are posted.",
         "In Teams: open the target channel → \"...\" next to the channel name → Get link to channel — the URL "
         "contains both GUIDs (groupId=<TEAMS_TEAM_ID> and the channel id segment is TEAMS_CHANNEL_ID, URL-decode "
         "the %3A back to ':'). Or via Graph Explorer: GET /me/joinedTeams, then GET /teams/{id}/channels. Docs: "
         "https://learn.microsoft.com/graph/teams-list-all-teams"),
    ],
    col_widths=[1.7, 2.3, 2.3],
)
doc.add_paragraph(
    "Every one of these is also printed inline by the corresponding step in deploy_solution.ps1 when you actually "
    "run it, so you do not need to keep this table open side-by-side — but this section is the single reference "
    "if you are filling in .env ahead of time."
)

add_heading(doc, "6.1 How to run the deployment script", level=2)
doc.add_paragraph(
    "Prerequisites: Azure CLI installed and logged in (az login, with the correct subscription selected via "
    "az account set), Python 3.11+ with pip, and an existing Fabric workspace + published Ontology data agent "
    "(Fabric IQ) and an existing Foundry project + deployed chat model + existing Foundry IQ knowledge agent — "
    "this accelerator reuses those, it does not build them from scratch."
)
add_code_block(doc, """
# From the repository root, in PowerShell:
cd C:\\path\\to\\Ins_FNOL_MicrosoftIQ_Logicapps_accelerator

# If this is the first time running scripts in this session, you may need
# to allow local script execution (one-time, per machine/user):
Set-ExecutionPolicy -Scope CurrentUser RemoteSigned

# Install Python dependencies used by the deployment/bootstrap scripts:
pip install -r requirements.txt

# Run the full deployment, step by step, pausing for manual/interactive
# actions where needed:
.\\deploy_solution.ps1

# If you are re-running after a partial deployment (e.g. the Agent Identity
# user and app registration already exist from a previous run), skip
# already-completed steps by number:
.\\deploy_solution.ps1 -SkipSteps 1,2
""")
doc.add_paragraph(
    "The script prints a clear 'STEP N' banner before each stage, and for the handful of steps that cannot be "
    "fully automated (assigning an M365 license, adding the user to Teams, authorizing the Office 365 connector in "
    "the Azure Portal, etc.) it prints the exact manual action required and waits for you to press Enter before "
    "continuing. It is safe to stop and re-run — most steps check for existing resources and skip recreation."
)

add_heading(doc, "6.2 What the script does, step by step", level=2)
doc.add_paragraph(
    "Each step below corresponds to a section in deploy_solution.ps1 and, if the step delegates to a standalone "
    "script, that script can also be run independently (useful when only one part of the environment needs to be "
    "redeployed)."
)

steps = [
    ("Step 0 — Prerequisites check",
     "Confirms az CLI login and Python are available; if .env does not exist yet, copies it from .env.example and "
     "opens it in Notepad for you to fill in before continuing.\n"
     "Docs: Azure CLI sign-in — https://learn.microsoft.com/cli/azure/authenticate-azure-cli\n"
     "Docs: Install Python — https://www.python.org/downloads/"),
    ("Step 1 — Create the Agent Identity user",
     "Prompts for the desired UPN, creates the user via 'az ad user create' (skips if it already exists), then "
     "pauses with instructions to complete four manual portal steps. Where to do each one:\n"
     "  1) Assign an M365 license including Exchange Online + Teams: https://admin.microsoft.com/Adminportal/"
     "Home#/users -> click the user -> Licenses and Apps -> select a license with Exchange Online + Teams -> "
     "Save changes.\n"
     "  2) Add the user as a Teams team MEMBER: https://admin.teams.microsoft.com/teams/manage-teams -> click "
     "the target team -> Members -> Add member -> search by UPN (or in the Teams app: ... next to the team "
     "name -> Add member).\n"
     "  3) Grant the user Contributor on the Fabric workspace hosting Fabric IQ: https://app.fabric.microsoft.com "
     "-> open the workspace -> Manage access -> Add people or groups -> search by UPN -> role 'Contributor' -> "
     "Add.\n"
     "  4) Grant the user the 'Foundry User' role on the Foundry account: https://portal.azure.com -> "
     "search for the Foundry/Cognitive Services account -> Access control (IAM) -> Add -> Add role assignment -> "
     "'Foundry User' -> Members -> select the user -> Review + assign. Cognitive Services User alone is not "
     "sufficient for the current Foundry data plane; keep/add it separately only if another component you use still "
     "depends on that legacy role.\n"
     "The script prints these same links inline when it pauses, so you do not need to keep this document open "
     "side-by-side while running it."),
    ("Step 2 — Register the Agent Identity public-client app",
     "Creates the 'FNOL-AgentIdentity-GraphPublicClient' app registration (skips if it already exists), adds the "
     "required delegated Graph (Mail.Read, Mail.Send, Sites.Read.All, Chat.Read, ChannelMessage.Send, "
     "ChannelMessage.Read.All), Power BI/Fabric "
     "(Item.Read.All, Item.Execute.All), and Azure SQL (user_impersonation) permissions, and requests admin "
     "consent. Prints the resulting appId and prompts you to confirm .env's AGENT_IDENTITY_PUBLIC_CLIENT_ID is set.\n"
     "Docs: App registrations — https://learn.microsoft.com/entra/identity-platform/quickstart-register-app\n"
     "Docs: Grant admin consent — https://learn.microsoft.com/entra/identity/enterprise-apps/grant-admin-consent\n"
     "Portal: https://portal.azure.com → Microsoft Entra ID → App registrations → "
     "FNOL-AgentIdentity-GraphPublicClient."),
    ("Step 3 — Bootstrap the Agent Identity's MSAL token cache",
     "Runs shared/bootstrap_agent_identity_tokens.py, which performs three interactive device-code sign-ins (as "
     "the Agent Identity user) for Graph, Fabric, and Azure SQL scopes, all persisted into one on-disk token cache. "
     "Also writes a base64-encoded copy for later use as chunked App Service app settings; the webapp reassembles "
     "AGENT_IDENTITY_TOKEN_CACHE_B64_0..N using AGENT_IDENTITY_TOKEN_CACHE_B64_COUNT at runtime and still supports "
     "the legacy single-setting name for backward compatibility.\n"
     "Docs: Device code flow — https://learn.microsoft.com/entra/identity-platform/v2-oauth2-device-code\n"
     "Sign in at https://microsoft.com/devicelogin, enter the code shown for each of the 3 prompts, and "
     "authenticate as the Agent Identity's UPN each time."),
    ("Step 4 — Create the CaseThreadMap table",
     "Runs shared/create_case_thread_map.py, which authenticates with the Agent Identity's SQL-scoped token and "
     "idempotently creates the CaseThreadMap table + unique index in the Fabric SQL Database.\n"
     "Docs: Fabric SQL Database connectivity — https://learn.microsoft.com/fabric/database/sql/connect"),
    ("Step 5 — Deploy the Work IQ webapp",
     "Runs triggers/webapp/deploy_webapp.py, which first checks whether the App Service Plan and Web App "
     "(WORKIQ_RESOURCE_GROUP / WORKIQ_WEBAPP_NAME in .env) already exist and auto-creates them if not (Linux/"
     "Python runtime, SKU from WORKIQ_SKU or B1 default, location from WORKIQ_LOCATION or westus3 default), then "
     "zips the webapp's Python files, deploys them via 'az webapp deploy', and pushes every required app setting "
     "(API key, Agent Identity client id, Fabric workspace/data agent ids, SQL connection string, and the base64 "
     "token cache from Step 3) directly from .env, so no secret is ever typed on the command line by hand. The "
     "token cache is pushed as chunked app settings (AGENT_IDENTITY_TOKEN_CACHE_B64_0..N plus "
     "AGENT_IDENTITY_TOKEN_CACHE_B64_COUNT) because a single App Service setting is too small for the full base64 "
     "blob.\n"
     "Docs: Create a Python App Service — https://learn.microsoft.com/azure/app-service/quickstart-python\n"
     "Docs: az webapp deploy (zip deploy) — https://learn.microsoft.com/cli/azure/webapp#az-webapp-deploy"),
    ("Step 6 — Create the Foundry connection for the webapp API key",
     "Runs foundry/create_workiq_connection.py, which creates (or updates) the 'workiq-api-key-conn' Custom Keys "
     "connection on the Foundry project via a direct ARM REST call (the SDK's ConnectionsOperations is read-only "
     "as of azure-ai-projects 2.5.0), so the orchestrator's OpenAPI tools can inject the x-api-key header without "
     "it ever appearing in the agent definition or source control.\n"
     "Docs: Foundry connections (Custom Keys) — https://learn.microsoft.com/azure/ai-foundry/how-to/connections-add"),
    ("Step 7 — Create the Foundry orchestrator agent",
     "Runs foundry/create_orchestrator_agent.py, which deletes any previously recorded orchestrator agent "
     "(orchestrator_agent_id.txt) and creates a fresh one wired to the Fabric IQ and Work IQ OpenAPI tools "
     "(via OpenApiConnectionAuthDetails against the connection from Step 6) and the existing Foundry IQ knowledge "
     "agent (via a ConnectedAgentTool). Records the new agent id in orchestrator_agent_id.txt.\n"
     "Docs: Azure AI Foundry Agent Service — https://learn.microsoft.com/azure/ai-services/agents/overview\n"
     "Docs: Connected agents / OpenAPI tools — "
     "https://learn.microsoft.com/azure/ai-services/agents/how-to/tools/openapi-spec"),
    ("Step 8 — Deploy the Logic Apps",
     "Runs triggers/logicapp/deploy_logic_apps.py, which now automates most of this step: it generates "
     "parameters.json automatically from .env values (foundryProjectEndpoint, teamId, channelId, "
     "fnolMailboxUserId, workiqBaseUrl, workiqApiKey) plus the orchestrator agent id recorded in "
     "foundry/orchestrator_agent_id.txt by Step 7 — no manual editing of that file is required. It then creates "
     "the Office 365 API Connection resource (idempotent) and writes connections.json automatically, and deploys "
     "both workflow.json and workflow-teams-reply-poller.json as two independent Consumption Logic App resources "
     "via direct ARM REST calls, printing each Logic App's managed-identity principalId for use in Step 9.\n"
     "The ONE step that genuinely cannot be automated is authorizing the Office 365 connection — this requires an "
     "interactive OAuth sign-in as the FNOL intake mailbox account, which only a human can complete in a browser. "
     "Do this at: Azure Portal → Resource groups → (your resource group) → the 'office365' connection resource → "
     "Edit API connection → Authorize → sign in as the FNOL intake mailbox → Save.\n"
     "Docs: Create a Consumption Logic App — "
     "https://learn.microsoft.com/azure/logic-apps/quickstart-create-example-consumption-workflow\n"
     "Docs: Office 365 Outlook connector authorization — https://learn.microsoft.com/connectors/office365/"),
    ("Step 9 — Grant the Logic Apps' Managed Identities their roles",
     "This step is now fully automated — no manual portal action or copy/pasted GUIDs needed. It looks up both "
     "Logic Apps' managed identity principal ids and the Foundry account's resource id directly (via 'az resource "
     "show' / 'az cognitiveservices account show', reading resource group/account names from .env), then: "
     "(1) grants BOTH Logic Apps' identities the 'Foundry User' Azure RBAC role scoped to the Foundry "
     "account via 'az role assignment create'; Cognitive Services User alone is insufficient for the current "
     "Foundry data plane, though both roles may coexist if other workloads still rely on the older role. "
     "(2) grants Logic App #2's identity the Microsoft Graph "
     "APPLICATION permission ChannelMessage.Read.All via a direct 'az rest' call to POST "
     "/servicePrincipals/{id}/appRoleAssignments — this is a Graph app-role grant, not an Azure RBAC role, so it "
     "cannot be done with 'az role assignment create', but 'az rest' (built into the Azure CLI you already use — "
     "no extra PowerShell module needed) handles it directly. Both grants are idempotent: re-running the step "
     "detects and skips ones already applied. If auto-discovery of any value fails (e.g. resource not found), the "
     "script falls back to prompting for it manually.\n"
     "Docs: Managed identity for Logic Apps — "
     "https://learn.microsoft.com/azure/logic-apps/authenticate-with-managed-identity\n"
     "Docs: Azure RBAC role assignment via CLI — "
     "https://learn.microsoft.com/azure/role-based-access-control/role-assignments-cli\n"
     "Docs: Grant a Graph app role to a managed identity — "
     "https://learn.microsoft.com/entra/identity/managed-identities-azure-resources/"
     "how-to-assign-app-role-managed-identity-powershell"),
]

for title, body in steps:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(body)

doc.add_page_break()

# ---------- 7. Verification checklist ----------
add_heading(doc, "7. Verification / Test Checklist", level=1)
add_table(doc,
    ["#", "Check", "Expected result", "Status (this deployment)"],
    [
        (1, "GET /healthz on the Work IQ webapp", "200 {\"status\":\"ok\"}", "PASS"),
        (2, "POST /query-ontology with a policy question", "200, grounded answer citing Fabric ontology data",
         "PASS (\"There are 59 open claims.\")"),
        (3, "POST /search-mail", "200, empty or matching results (no auth error)", "PASS"),
        (4, "POST /search-documents", "200, empty or matching results (no auth error)", "PASS"),
        (5, "POST /send-email", "200 when explicitly invoked with valid to/subject/bodyText", "PASS"),
        (6, "POST /post-new-case-alert", "200, new Teams root message id returned", "PASS"),
        (7, "POST /insert-case-thread-map then duplicate insert", "First succeeds; duplicate rejected (PK)", "PASS"),
        (8, "POST /post-case-reply nested under a root message", "200, new reply message id, nested in Teams UI", "PASS"),
        (9, "POST /resolve-case-reply", "200, correct CaseId + foundry_thread_id returned", "PASS"),
        (10, "Create orchestrator agent (create_orchestrator_agent.py)", "Agent created with 3 tools wired", "PASS"),
        (11, "Two concurrent end-to-end case runs (automated test)", "Distinct Foundry thread ids AND distinct "
         "Teams root message ids; each case's answer reflects its own policy data with no cross-contamination",
         "PASS"),
    ],
    col_widths=[0.4, 2.4, 2.6, 1.6],
)

doc.add_page_break()

# ---------- 9. End-to-end manual test ----------
add_heading(doc, "9. End-to-End Manual Test", level=1)
doc.add_paragraph(
    "This section is a simple, step-by-step manual test anyone can follow after deployment (Steps 0-9 in Section "
    "6) to confirm the whole pipeline works for real, from an inbound FNOL email through to a Teams reply. No "
    "scripting or coding required — just an email client, Microsoft Teams, and a few minutes of waiting between "
    "steps for the Logic Apps' 1-minute polling interval."
)
add_callout(
    doc, "Before you start",
    "Confirm all of the following, or the test will fail partway through: (1) Steps 0-9 of the deployment "
    "runbook (Section 6) have completed successfully; (2) the Office 365 connection was authorized (Section 6.2, "
    "Step 8's one manual action); (3) the Fabric capacity backing Fabric IQ / CaseThreadMap is Active, not paused; "
    "(4) you have access to send email to the Agent Identity's mailbox (its UPN is in your .env's "
    "AGENT_IDENTITY_UPN) and are a member of the Teams channel configured in TEAMS_TEAM_ID/TEAMS_CHANNEL_ID.",
)

add_heading(doc, "Step 1 — Send a test FNOL email", level=2)
doc.add_paragraph(
    "From any mailbox you have access to (your own is fine — it does not have to be a special test account), "
    "send a new email:"
)
add_table(
    doc,
    ["Field", "Value"],
    [
        ("To", "The Agent Identity's UPN (e.g. svc-fnol-agent@yourtenant.onmicrosoft.com — this is the FNOL "
                "intake mailbox; find it in .env's AGENT_IDENTITY_UPN)"),
        ("Subject", "TEST FNOL - Auto Accident Claim"),
        ("Body", "Any short claim-like description works, e.g.: \"Policy number POL-00005. Rear-end collision "
                  "at a stop light on Main St, no injuries reported. 2022 Sedan.\""),
    ],
    col_widths=[1.2, 5.3],
)
doc.add_paragraph("Send the email, then note the time — the next step can take up to a minute to pick it up.")

add_heading(doc, "Step 2 — Confirm the email-intake Logic App ran", level=2)
add_numbered(doc, [
    "Go to https://portal.azure.com → Resource groups → (your LOGICAPP_RESOURCE_GROUP, e.g. RG_AutoFNOL_"
    "Logicapps) → open the Logic App named in .env's LOGICAPP_EMAIL_INTAKE_NAME (default: la-fnol-email-intake).",
    "In the left navigation, click Overview, then look at the \"Runs history\" list at the bottom of the page.",
    "Within about a minute of sending the email, a new run should appear. Click it to open the run details.",
    "Expected result: every step shows a green checkmark (Succeeded) — in particular \"When_a_new_email_"
    "arrives\", \"Create_Foundry_thread\", \"Start_run\", \"Until_run_completes\", \"Post_new_case_alert_to_"
    "Teams\", and \"Persist_CaseThreadMap_row\".",
])
add_callout(
    doc, "If no run appears after 2-3 minutes",
    "Check: the Office 365 connection is authorized (Portal → same resource group → the 'office365' connection "
    "→ status should not say \"Error\" or \"Unauthenticated\"); the test email actually landed in the Agent "
    "Identity's Inbox (not spam/another folder — the trigger only watches the Inbox); and the Logic App is "
    "enabled (Overview page shows a \"Disable\" button, not \"Enable\").",
    color_hex="FCE4E4",
)

add_heading(doc, "Step 3 — Confirm the Teams alert was posted", level=2)
add_numbered(doc, [
    "Open Microsoft Teams and go to the channel configured in .env's TEAMS_TEAM_ID/TEAMS_CHANNEL_ID.",
    "Look for a new top-level message (not a reply) summarizing the test claim — this is the orchestrator "
    "agent's triage output, and it should reference the policy/claim details you put in the email body.",
    "This message is the new case's \"root\" message — every future reply for this exact case will be nested "
    "under it, so you can find the whole conversation for this case in one place no matter how many other cases "
    "are posted around the same time.",
])

add_heading(doc, "Step 4 — Reply in Teams to test the follow-up path", level=2)
add_numbered(doc, [
    "In Teams, click Reply directly under the new case's message from Step 3 (not a new message — an actual "
    "threaded reply).",
    "Type a short follow-up, e.g.: \"What is the current claim status for this policy?\" and send it.",
    "Wait up to 30 seconds for the reply-poller Logic App's next polling cycle (LOGICAPP_REPLY_POLLER_NAME, "
    "default: la-fnol-teams-reply-poller) to pick it up.",
    "Refresh the Teams channel — a new reply from the orchestrator agent should appear nested under the SAME "
    "case thread, answering your follow-up question using the same case's context (it should not ask you to "
    "re-explain the claim from scratch).",
])
doc.add_paragraph(
    "To confirm this at the platform level instead of just visually in Teams: Portal → your LOGICAPP_RESOURCE_"
    "GROUP → open la-fnol-teams-reply-poller → Overview → Runs history — a new run should appear shortly after "
    "your reply, with all steps (List_channel_messages, Filter_to_new_human_replies, For_each_new_reply, "
    "Resolve_case_from_reply, Post_reply_to_Foundry_thread, Post_ack_reply_to_Teams) shown as Succeeded."
)

add_heading(doc, "Step 5 — (Optional) Confirm the CaseThreadMap row directly", level=2)
doc.add_paragraph(
    "For a deeper check that doesn't rely on the Teams UI, you can query the CaseThreadMap table directly to see "
    "the row this test case created:"
)
add_code_block(doc, """
# From the repository root, in PowerShell (requires the Agent Identity's
# token cache to already be bootstrapped - see Section 6.2, Step 3):
python -c "
import sys, struct
sys.path.insert(0, 'shared')
from agent_identity_auth import get_agent_token
import pyodbc, os
from dotenv import load_dotenv
load_dotenv()
token = get_agent_token(['https://database.windows.net/.default'])
tb = token.encode('utf-16-le')
ts = struct.pack('=i', len(tb)) + tb
conn = pyodbc.connect(os.environ['CASETHREADMAP_FABRIC_SQL_CONNECTION_STRING'], attrs_before={1256: ts})
cur = conn.cursor()
cur.execute('SELECT TOP 5 * FROM CaseThreadMap ORDER BY CreatedUtc DESC')
for row in cur.fetchall(): print(row)
"
""")
doc.add_paragraph(
    "Expected result: a row whose CreatedUtc is around the time you sent the test email, with a non-empty "
    "TeamId, ChannelId, RootMessageId (matching the Teams message from Step 3), and FoundryThreadId."
)

add_heading(doc, "What a fully passing test proves", level=2)
add_bullets(doc, [
    "The Office 365 email trigger, Foundry thread creation, orchestrator agent run, Teams posting, and "
    "CaseThreadMap persistence all worked together for a real inbound case (Logic App #1's entire pipeline).",
    "The Teams reply-poller correctly detected a genuine human reply, resolved it back to the right case via "
    "CaseThreadMap, continued the SAME Foundry thread (not a new one), and posted the follow-up answer nested "
    "under the correct case (Logic App #2's entire pipeline).",
    "Every Agent Identity-authenticated call across the whole solution — Graph mail read, Graph Teams posting, "
    "Fabric ontology queries (if your test question triggered one), and Fabric SQL reads/writes — is working "
    "end to end with real, live credentials.",
])
doc.add_paragraph(
    "If you want to test with multiple cases running concurrently (to verify the isolation guarantees from "
    "Section 5), simply repeat Step 1 with 2-3 different test emails a few seconds apart, then confirm in Step 3 "
    "that each produces its OWN separate top-level Teams message with no shared/mixed content."
)

doc.add_page_break()

# ---------- 10. Operational notes ----------
add_heading(doc, "10. Operational Notes & Known Limitations", level=1)
add_bullets(doc, [
    "Fabric capacity must be Active (not paused) for BOTH Fabric IQ ontology queries and CaseThreadMap SQL reads/"
    "writes to succeed — pausing it to save cost between testing sessions is fine, but it must be resumed before "
    "any end-to-end run and takes a few minutes to fully warm up after resume (initial SQL connections may see a "
    "transient login timeout during warm-up).",
    "The Agent Identity's MSAL token cache is bootstrapped once via interactive device-code sign-in per resource "
    "scope (Graph, Fabric, SQL) and then silently refreshes — no further interactive sign-in should be required "
    "unless the refresh token itself expires/is revoked (tenant conditional access policy dependent) or new scopes "
    "are added to the app registration later. In App Service, the seed cache is stored as chunked settings "
    "(AGENT_IDENTITY_TOKEN_CACHE_B64_0..N plus _COUNT) and reassembled at runtime; the legacy single-setting "
    "AGENT_IDENTITY_TOKEN_CACHE_B64 is still honored for compatibility.",
    "There are two copies of agent_identity_auth.py (shared/ and triggers/webapp/) and no automated sync between "
    "them. If you fix a bug only in shared/, the deployed webapp will NOT pick it up until triggers/webapp/"
    "agent_identity_auth.py is updated and the webapp is redeployed.",
    "The reply-poller Logic App (workflow-teams-reply-poller.json) uses simple polling rather than a Bot Framework "
    "webhook, trading a small latency (up to the 30-second recurrence interval) for zero additional bot "
    "infrastructure. Its trigger also forces runtimeConfiguration.concurrency.runs = 1 so overlapping poll cycles "
    "queue instead of racing to post another message into the same Foundry thread while a prior run is active, "
    "which avoids intermittent 400 'Can't add messages to thread while a run is active' failures. Swap in Graph "
    "delta queries (deltaLink) in production to avoid re-scanning the whole channel on every poll.",
    "Teams reply/alert HTML is post-processed with inline block-level margin styling and urgency-line colorization "
    "before posting because Teams ignores normal browser default margins on block elements; without those helpers, "
    "multi-section agent replies render as one dense block with weak urgency emphasis.",
    "All secrets (API key, connection strings, client ids treated as sensitive, token cache) are kept out of source "
    "control via .env/.gitignore and Azure App Service application settings — never hardcoded in any .py/.json file "
    "committed to the repository.",
    "This solution is intentionally READ-ONLY against Fabric IQ's Policy/Claim data — the orchestrator's "
    "instructions explicitly prohibit creating or modifying claim/policy records.",
    "As of this document's date, all work has been deployed and verified in the live Azure/Entra tenant but has "
    "NOT been committed or pushed to the GitHub repository, per explicit instruction — treat the working tree as "
    "the source of truth until an explicit commit/push is requested.",
])

import os
_out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "End_to_End_Deployment_Guide.docx")
doc.save(_out_path)
print(f"Saved {_out_path}")
