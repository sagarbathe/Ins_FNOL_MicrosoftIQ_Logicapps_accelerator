# Generates docs/OBO_Bot_Teams_Design.docx
# Run: python generate_obo_bot_design_doc.py

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        doc.add_paragraph(item, style=style)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    # light gray background via shading on the paragraph's containing "cell"-like box
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


doc = Document()

# ---------- Title ----------
title = doc.add_heading("AutoFNOL Accelerator: Delegated Teams Authentication via Bot Framework + On-Behalf-Of (OBO) Flow", level=0)
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run(
    "Design for fully-automated, delegated Fabric Data Agent (Ontology) access from the Foundry orchestrator, "
    "using a Teams Bot's OAuth Token Store to eliminate the service-principal ontology-auth gap."
)
subtitle_run.italic = True
subtitle_run.font.size = Pt(12)

meta = doc.add_paragraph()
meta.add_run("Tenant: ").bold = True
meta.add_run("mngenvmcap146722.onmicrosoft.com\n")
meta.add_run("Primary user identity: ").bold = True
meta.add_run("sagarbathe@mngenvmcap146722.onmicrosoft.com\n")
meta.add_run("Document date: ").bold = True
meta.add_run("2026-08-30")

doc.add_page_break()

# ---------- 1. Problem statement ----------
add_heading(doc, "1. Problem Statement", level=1)
doc.add_paragraph(
    "The Foundry orchestrator's Fabric IQ tool must query a Fabric Data Agent backed by an Ontology data source. "
    "Testing (see companion diagnostic notebooks) conclusively proved:"
)
add_bullets(doc, [
    "Service-principal (app-only) tokens successfully authenticate at the MCP transport layer (handshake, tool "
    "discovery, and tool invocation all succeed).",
    "However, the Data Agent's internal query into the Ontology/Semantic Model engine fails authorization when "
    "the underlying token represents a pure SP identity - the agent's response explicitly states an "
    "\"authorization error\" prevents it from looking up ontology data.",
    "The identical call, using a delegated user token (notebookutils.credentials / a signed-in user), succeeds "
    "and returns correct ontology-backed answers.",
])
doc.add_paragraph(
    "Conclusion: Fabric ontology-backed Data Agents currently require a delegated (user-context) token for the "
    "actual data query, even though the outer MCP transport accepts SP tokens. Because the Data Agent's source "
    "must remain the Ontology (not a Lakehouse/Warehouse), and the orchestrator must run fully unattended "
    "(triggered by an inbound email, with no interactive user in the loop), a mechanism is needed to obtain a "
    "valid delegated token with zero ongoing human interaction."
)

add_heading(doc, "1.1 Why On-Behalf-Of (OBO) Alone Isn't Sufficient", level=2)
doc.add_paragraph(
    "OBO exchanges an existing delegated (user) token for a new token scoped to a different resource - it cannot "
    "manufacture a first delegated token from nothing. Because the email trigger has no signed-in user anywhere "
    "in the chain, there is initially no token to exchange. The missing piece is a way to silently obtain a cached "
    "delegated token for a real, designated user - "
    "sagarbathe@mngenvmcap146722.onmicrosoft.com - without that user interacting each time. Azure Bot Service's "
    "OAuth Token Store solves exactly this problem, and OBO is then used on top of it to mint the Fabric-scoped token."
)

# ---------- 2. Solution overview ----------
add_heading(doc, "2. Solution Overview", level=1)
doc.add_paragraph(
    "Introduce a Teams Bot (Azure Bot Service, a pure Azure resource - no Power Automate or Copilot Studio) as the "
    "component that owns the Teams conversation. The bot's built-in OAuth Token Store performs a one-time "
    "interactive sign-in (during setup only) for the designated user, then silently refreshes and serves that "
    "user's delegated token forever after, with no further prompts. The orchestrator uses that cached token as "
    "input to an OBO exchange to mint a Fabric-scoped delegated token, which the Fabric IQ tool then uses to query "
    "the ontology successfully."
)

add_heading(doc, "2.1 New / Changed Components vs. Current Architecture", level=2)
add_table(
    doc,
    ["Component", "Type", "Role", "Status"],
    [
        ("Azure Bot Service resource", "Microsoft.BotService/botServices", "Registers the bot identity in Entra; enables the Teams channel; owns the OAuth Connection Setting and Token Store.", "New"),
        ("Bot backend (Bot Framework SDK)", "Azure Function or App Service", "Hosts the bot's message-handling logic: receives/sends Teams messages, calls the User Token Service, triggers the orchestrator.", "New (can reuse existing app-autofnol-workiq App Service Plan or a new small App Service)"),
        ("OAuth Connection Setting", "Bot resource configuration", "Entra app registration + Fabric API scope (https://api.fabric.microsoft.com/.default), configured once on the Bot resource.", "New"),
        ("Foundry orchestrator (Foundry Agent Service)", "Existing", "Receives the delegated token (via the bot), performs the OBO exchange, and invokes Fabric IQ / Work IQ tools.", "Modified: add OBO exchange step before calling Fabric IQ"),
        ("Logic App (email trigger)", "Existing", "Triggers on inbound email to sagarbathe@mngenvmcap146722.onmicrosoft.com, extracts FNOL details, calls the orchestrator.", "Modified: now calls the bot (or a shared Cosmos/queue) to route the reply to the correct Teams conversation instead of posting directly via Graph HTTP"),
        ("CaseThreadMap (Fabric SQL Database)", "Existing", "Maps each FNOL case to its Teams conversation/thread reference, preserving context across concurrent emails.", "Unchanged"),
        ("Work IQ (App Service)", "Existing", "Graph search wrapper (SharePoint documents + mailbox search).", "Unchanged"),
        ("Fabric IQ (Data Agent, MCP)", "Existing", "Ontology-backed Fabric Data Agent, queried via MCP.", "Unchanged (data source stays Ontology)"),
    ],
    col_widths=[1.6, 1.3, 2.6, 1.5],
)

# ---------- 3. One-time setup (interactive, done once) ----------
add_heading(doc, "3. One-Time Setup (Interactive, Done Once - Not Per-Email)", level=1)
doc.add_paragraph(
    "This step happens once during deployment (or rarely, if the cached token fully expires after prolonged "
    "inactivity or a credential reset). It is the only human touchpoint in the entire design."
)
add_bullets(doc, [
    "Deploy the Azure Bot Service resource and register the Teams channel.",
    "Configure an OAuth Connection Setting on the bot pointing to an Entra app registration, requesting the "
    "https://api.fabric.microsoft.com/.default scope (delegated).",
    "In a Teams chat with the bot, sagarbathe@mngenvmcap146722.onmicrosoft.com completes a single OAuthPrompt "
    "sign-in/consent flow.",
    "Bot Framework's Token Store persists a refresh token for this user, keyed by (userId, channelId, "
    "connectionName), in the Bot Service's managed, encrypted token store.",
], style="List Number")
doc.add_paragraph(
    "From this point forward, the token is retrieved and refreshed silently by the bot backend on every "
    "subsequent request - no further prompts, ever, unless the refresh token itself is revoked or expires from "
    "prolonged disuse."
)

# ---------- 4. End-to-end workflow ----------
add_heading(doc, "4. End-to-End Workflow (Fully Automated, Per Email)", level=1)

steps = [
    ("1. Email arrives", "An FNOL email arrives at sagarbathe@mngenvmcap146722.onmicrosoft.com."),
    ("2. Logic App trigger", "The existing Logic App's inbound-email trigger fires, parses sender/subject/body/attachments, and generates a unique CaseId."),
    ("3. CaseThreadMap insert", "The Logic App inserts a new row into the Fabric SQL Database CaseThreadMap table: (CaseId, EmailMessageId, TeamsConversationId = NULL initially, Status = 'New', CreatedAt)."),
    ("4. Orchestrator invocation", "The Logic App calls the Foundry orchestrator agent (Foundry Agent Service), passing the CaseId and email content."),
    ("5. Silent delegated token retrieval", "Before calling Fabric IQ, the orchestrator (via the bot backend, which it calls as a lightweight internal API) requests: UserTokenClient.GetUserTokenAsync(sagarbathe_aad_object_id, connectionName, channelId). This returns a valid delegated access token silently, refreshing under the hood if needed - no prompt."),
    ("6. OBO exchange", "The orchestrator's confidential client application calls acquire_token_on_behalf_of(user_assertion=<token from step 5>, scopes=['https://api.fabric.microsoft.com/.default']). This mints a new delegated, Fabric-scoped access token, still carrying the user's identity and permissions."),
    ("7. Fabric IQ (Ontology) query", "The orchestrator calls the Fabric Data Agent's MCP endpoint using the OBO-minted delegated token. The ontology query now succeeds (delegated auth is honored), returning real policy/coverage/claims data."),
    ("8. Work IQ query (as needed)", "In parallel or as a follow-up tool call, the orchestrator calls Work IQ (App Service, SP auth - unaffected by this change) to search SharePoint documents or mailbox content."),
    ("9. Response composed", "The orchestrator composes a response (e.g., coverage summary, missing-info questions) grounded in the ontology + Work IQ results."),
    ("10. Teams message posted via bot", "The orchestrator hands the response back to the bot backend, which posts it into the Teams conversation for this case. If this is the first message for the case, the bot starts a new 1:1 or channel conversation and captures the resulting TeamsConversationId / activity/thread reference."),
    ("11. CaseThreadMap updated", "The bot backend (or orchestrator) updates the CaseThreadMap row for this CaseId with the TeamsConversationId and the Bot Framework conversation reference (serialized ConversationReference), setting Status = 'Awaiting Reply'."),
    ("12. User replies in Teams", "sagarbathe (or an adjuster) replies inside that specific Teams conversation thread - because each case has its own distinct conversation/thread (see Section 5), there is no ambiguity about which case the reply belongs to."),
    ("13. Bot receives reply, resolves CaseId", "The bot's message handler receives the incoming Teams activity, looks up CaseThreadMap by TeamsConversationId to resolve the CaseId, then re-invokes the orchestrator with that CaseId and the new message content - repeating steps 5-11 for the follow-up turn."),
]

for title_text, body_text in steps:
    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph(body_text)

# ---------- 5. Concurrency / thread isolation ----------
add_heading(doc, "5. Preserving Context Across Concurrent Emails (Thread Isolation)", level=1)
doc.add_paragraph(
    "The Bot Framework model naturally solves the multi-email/multi-conversation problem that was previously "
    "ambiguous with a single shared Teams channel:"
)
add_bullets(doc, [
    "Each CaseId gets its own dedicated Teams conversation (e.g., the bot starts a new 1:1 chat, or a new thread "
    "within a channel, per case) rather than posting all cases into one shared channel feed.",
    "The mapping CaseId <-> TeamsConversationId <-> ConversationReference is persisted in CaseThreadMap "
    "(Fabric SQL Database), which already exists in the current architecture for exactly this purpose.",
    "Because Bot Framework activities always carry the conversation's own ID, an incoming reply is unambiguously "
    "attributable to its case - the end user simply keeps replying inside that case's own conversation with the "
    "bot; there is no need for the user to remember or select the correct thread manually.",
    "This directly resolves the earlier concern (\"multiple emails arrive at once, how does the user know which "
    "Teams message to reply to\") - because each case IS its own conversation, not a subset of a single shared "
    "chat feed.",
])

# ---------- 6. Data flow diagram (text) ----------
add_heading(doc, "6. Data Flow Summary", level=1)
add_code_block(doc, (
    "Inbound Email\n"
    "     |\n"
    "     v\n"
    "Logic App (email trigger)\n"
    "     |-- INSERT CaseThreadMap (CaseId, Status='New')\n"
    "     v\n"
    "Foundry Orchestrator Agent\n"
    "     |-- calls Bot backend: GetUserTokenAsync(sagarbathe, connectionName)  --> delegated user token (cached, silent)\n"
    "     |-- OBO exchange: acquire_token_on_behalf_of(user token, fabric scope) --> delegated Fabric token\n"
    "     |-- Fabric IQ (MCP) query using delegated Fabric token --> Ontology data\n"
    "     |-- Work IQ (App Service, SP auth) query --> SharePoint / mailbox data\n"
    "     v\n"
    "Response composed\n"
    "     v\n"
    "Bot backend posts/starts Teams conversation\n"
    "     |-- UPDATE CaseThreadMap (TeamsConversationId, ConversationReference, Status='Awaiting Reply')\n"
    "     v\n"
    "Teams (dedicated conversation per case)\n"
    "     |\n"
    "     v  (user reply)\n"
    "Bot backend receives activity\n"
    "     |-- SELECT CaseThreadMap WHERE TeamsConversationId = <incoming> --> CaseId\n"
    "     v\n"
    "Foundry Orchestrator Agent (re-invoked with resolved CaseId + new message)\n"
    "     ... (repeat token/OBO/Fabric IQ/Work IQ steps) ...\n"
))

# ---------- 7. Azure resources to deploy ----------
add_heading(doc, "7. Azure Resources to Deploy", level=1)
add_table(
    doc,
    ["Resource", "Resource Type", "Resource Group", "Notes"],
    [
        ("bot-autofnol-teams", "Microsoft.BotService/botServices", "RG_AutoFNOL_Logicapps", "Registers Entra app for the bot identity; enable Teams channel."),
        ("app-autofnol-botbackend", "App Service (Linux, Python) or reuse app-autofnol-workiq", "RG_AutoFNOL_Logicapps", "Hosts Bot Framework SDK message handler + Token Store calls + OBO exchange logic."),
        ("OAuth Connection Setting", "Bot resource sub-configuration", "(on bot-autofnol-teams)", "Entra app registration + https://api.fabric.microsoft.com/.default scope."),
        ("AutoFNOL-Logicapps-GraphApp (existing)", "Entra App Registration", "N/A (Entra)", "Reused; may need 'Allow public client flows' and delegated Fabric API permission added for OBO."),
    ],
    col_widths=[2.0, 2.4, 1.6, 2.5],
)

doc.add_paragraph(
    "No VNets or Private Endpoints are required anywhere in this design, consistent with the existing constraint. "
    "No Power Automate or Copilot Studio components are introduced - the Bot Service is a pure Azure/Bot Framework "
    "SDK resource."
)

# ---------- 8. Security / operational notes ----------
add_heading(doc, "8. Security and Operational Notes", level=1)
add_bullets(doc, [
    "The delegated token always represents sagarbathe's own permissions - access to ontology data is bounded by "
    "what this user is actually authorized to see, preserving a meaningful audit trail (versus a generic service "
    "account or SP).",
    "The Bot Framework Token Store encrypts and manages refresh-token storage; no raw secrets are handled by "
    "custom code beyond calling the SDK's UserTokenClient.",
    "If the cached refresh token is ever revoked (password change, admin revocation, or long inactivity), the "
    "bot backend will receive an auth failure from GetUserTokenAsync and should raise an alert; a repeat of the "
    "one-time interactive sign-in (Section 3) resolves this.",
    "This pattern currently depends on one designated real user's delegated permissions; if multiple adjusters "
    "need distinct ontology-access scopes, consider provisioning one dedicated 'automation' user per team/queue "
    "rather than a personal account, while keeping the same Bot Token Store + OBO mechanism.",
])

doc.save(r"C:\Sagar\MicrosoftIQ\Ins_FNOL_MicrosoftIQ_Logicapps_accelerator\docs\OBO_Bot_Teams_Design.docx")
print("Saved docx")
