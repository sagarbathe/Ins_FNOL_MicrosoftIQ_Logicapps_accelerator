"""
Generates a Word document summarizing:
1. Solution design (architecture, components, dataflow)
2. Prerequisites required to deploy
3. Deployment steps taken, with status (Succeeded / Failed / Not yet run)

Run: python generate_status_doc.py
Output: FNOL_Logicapps_Solution_Status.docx (in this folder)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- helpers ----------
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_status_table(rows, headers=("Component / Step", "Status", "Details")):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    status_colors = {
        "Succeeded": "C6EFCE",
        "Passed": "C6EFCE",
        "Failed": "FFC7CE",
        "Blocked": "FFC7CE",
        "Not yet run": "FFEB9C",
        "Partial": "FFEB9C",
        "N/A": "D9D9D9",
    }
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
        status_val = row[1]
        set_cell_shading(cells[1], status_colors.get(status_val, "FFFFFF"))
    doc.add_paragraph()

# ==================== TITLE PAGE ====================
title = doc.add_heading('Auto FNOL Triage \u2013 Foundry & Logic Apps Accelerator', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Solution Design, Prerequisites, and Deployment Status Report')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
sub.runs[0].font.size = Pt(14)
meta = doc.add_paragraph('Tenant: mngenvmcap146722.onmicrosoft.com   |   Subscription: ME-MngEnvMCAP146722-sagarbathe-1')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
doc.add_page_break()

# ==================== 1. OVERVIEW ====================
add_heading('1. Solution Overview', level=1)
doc.add_paragraph(
    "This solution automates First Notice of Loss (FNOL) triage for auto insurance claims. "
    "It replaces the Copilot Studio + Power Automate orchestration layer used in a prior accelerator "
    "with a fully Azure-native stack: an Azure AI Foundry Agent Service orchestrator, a Logic App "
    "as the event trigger and integration glue, a Fabric SQL Database for Teams-thread continuity "
    "state, and a lightweight Azure App Service acting as a Microsoft Graph tool bridge (\"Work IQ\"). "
    "The solution intentionally avoids any VNets or Private Endpoints, per explicit design requirement."
)
doc.add_paragraph(
    "Three existing, reused building blocks are intentionally kept unchanged from prior FNOL "
    "accelerator work in this tenant:"
)
for b in [
    "Fabric IQ \u2013 an Ontology + Fabric Data Agent published over a Lakehouse containing "
    "Policy, Vehicle, Claim, Adjuster, RepairShop, FraudSignal, and SubrogationFlag data.",
    "Foundry IQ \u2013 an existing Azure AI Search-grounded Foundry knowledge agent "
    "(agent id asst_thLtP20E4lDRqxiaW3rtftLF) used as a connected sub-agent for policy/claims documents.",
    "Work IQ \u2013 a Microsoft Graph-based tool for searching SharePoint/OneDrive documents and Outlook "
    "mail for fresh, ungoverned operational content not present in the governed knowledge base.",
]:
    doc.add_paragraph(b, style='List Bullet')

# ==================== 2. ARCHITECTURE ====================
add_heading('2. Architecture & Components', level=1)

add_heading('2.1 High-Level Dataflow', level=2)
for step in [
    "1. An email describing a new loss event arrives in a monitored mailbox.",
    "2. A Logic App (Office 365 Outlook trigger) fires on new mail, extracts the email body/attachments, "
    "and calls the Foundry orchestrator agent with the claim narrative.",
    "3. The Foundry orchestrator agent (Azure AI Foundry Agent Service, Assistants-style) reasons over the "
    "request and calls out to three tools as needed:",
    "     a. Foundry IQ (connected agent) \u2013 for governed policy/claims knowledge.",
    "     b. Fabric IQ (Fabric Data Agent MCP endpoint) \u2013 for live ontology-grounded structured data "
    "(coverage limits, policy status, vehicle/adjuster assignment, fraud/subrogation flags).",
    "     c. Work IQ (Azure App Service / Graph API wrapper) \u2013 for ad hoc SharePoint/OneDrive document "
    "search and Outlook mail search not covered by the governed knowledge base.",
    "4. The orchestrator's triage recommendation and case summary are posted to a Microsoft Teams channel "
    "message via the Logic App's Teams action.",
    "5. A CaseThreadMap row is written to the Fabric SQL Database, mapping the case/claim ID to the Teams "
    "message/thread ID, so that follow-up replies in Teams can be correlated back to the same case even "
    "when multiple emails/cases are being processed concurrently.",
    "6. Adjusters/agents can reply in the Teams thread; a reply-triggered flow (not yet built) would use "
    "CaseThreadMap to resume the correct case context in a follow-up call to the orchestrator.",
]:
    doc.add_paragraph(step)

add_heading('2.2 Components Table', level=2)
comp_rows = [
    ("Azure AI Foundry Project", "sbazureaimodels-project01 (existing Foundry account sbazureaimodels, resource group RG_openai)", ""),
    ("Foundry Orchestrator Agent", "New \u2013 Assistants-style agent (azure-ai-agents SDK) with 3 tools: Foundry IQ (connected agent), Fabric IQ (OpenAPI facade), Work IQ (OpenAPI facade)", ""),
    ("Foundry IQ Knowledge Agent", "Reused \u2013 asst_thLtP20E4lDRqxiaW3rtftLF (Azure AI Search-grounded)", ""),
    ("Fabric IQ Data Agent", "Reused \u2013 DA_AutoFNOL_Ontology, workspace WS_AutoFNOL, over ontology AutoFNOL_Ontology & Lakehouse LH_AutoFNOL", ""),
    ("Work IQ Web Service", "New \u2013 Azure App Service (Linux, Python, B1) exposing /search-documents, /search-mail, /query-ontology, /healthz", ""),
    ("Fabric SQL Database (CaseThreadMap)", "New \u2013 AutoFNOLCaseThreadMap in workspace WS_AutoFNOL; stores CaseId \u2194 Teams channel/message/thread ID mapping", ""),
    ("Logic App", "New \u2013 Office 365 Outlook trigger \u2192 Foundry orchestrator call \u2192 Teams post \u2192 CaseThreadMap upsert", ""),
    ("Entra App Registration (Graph)", "New \u2013 AutoFNOL-Logicapps-GraphApp; app-only Graph permissions Sites.Read.All, Mail.Read; used by Work IQ and by Fabric IQ facade for Fabric API auth", ""),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Component"; hdr[1].text = "Description"; hdr[2].text = ""
for c in [hdr[0], hdr[1]]:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
table.columns[2].width = Inches(0.1)
for name, desc, _ in comp_rows:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc
doc.add_paragraph()

add_heading('2.3 Design Constraints', level=2)
for c in [
    "No VNets and no Private Endpoints anywhere in the solution (explicit requirement) \u2014 all Azure "
    "resources use public HTTPS endpoints, secured via Entra ID auth, API keys, and/or IP-agnostic RBAC "
    "rather than network isolation.",
    "No Copilot Studio and no Power Automate \u2014 replaced by Azure AI Foundry Agent Service (orchestrator) "
    "and Azure Logic Apps (integration/event flow).",
    "Fabric SQL Database used instead of Azure SQL Database for the CaseThreadMap store, per explicit "
    "preference to keep case-thread state inside the Fabric estate alongside the existing Fabric IQ Lakehouse.",
    "This document and the underlying repository intentionally do not reference the original Copilot "
    "Studio/Power Automate accelerator repository; only this solution's own architecture is described.",
]:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ==================== 3. PREREQUISITES ====================
add_heading('3. Prerequisites', level=1)
for c in [
    "An Azure subscription with Contributor access to create resource groups, App Services, and Entra "
    "app registrations (used: ME-MngEnvMCAP146722-sagarbathe-1).",
    "An existing Azure AI Foundry account/project with at least one chat-capable model deployment "
    "(used: sbazureaimodels-project01, model gpt-4o).",
    "An existing Microsoft Fabric workspace on a paid F2+ capacity, containing:",
    "     \u2013 A Lakehouse with Policy/Claim/Vehicle/Adjuster/RepairShop tables populated with sample data.",
    "     \u2013 An Ontology item mapping those tables to business entities.",
    "     \u2013 A published Fabric Data Agent over that ontology (used: DA_AutoFNOL_Ontology on workspace "
    "WS_AutoFNOL, capacity sbfabriccapeastus2, SKU F16, region East US 2).",
    "An existing Azure AI Search index + Foundry knowledge agent for governed policy/claims documents "
    "(reused from a prior accelerator, agent id asst_thLtP20E4lDRqxiaW3rtftLF).",
    "Microsoft Entra ID permissions to create an app registration and grant Microsoft Graph application "
    "permissions (Sites.Read.All, Mail.Read) with admin consent.",
    "A Microsoft 365 mailbox to monitor for incoming FNOL emails, and a Teams team/channel to post case "
    "alerts to.",
    "Python 3.11+ locally (for the Foundry agent creation script and Fabric SQL schema push), with "
    "packages: azure-ai-projects, azure-ai-agents, azure-identity, pyodbc, msal, mcp, flask, gunicorn.",
    "ODBC Driver 18 for SQL Server (for applying the CaseThreadMap schema to the Fabric SQL Database "
    "from a local machine).",
]:
    style = 'List Bullet' if not c.startswith('     ') else None
    doc.add_paragraph(c, style=style)

doc.add_page_break()

# ==================== 4. DEPLOYMENT STEPS & STATUS ====================
add_heading('4. Deployment Steps Taken & Status', level=1)
doc.add_paragraph(
    "The table below lists every deployment step attempted in this session, in order, with its outcome. "
    "\u201cSucceeded\u201d steps are verified working; \u201cFailed/Blocked\u201d steps have a known root cause "
    "and remediation path; \u201cNot yet run\u201d steps have not been attempted yet."
)

add_heading('4.1 Resource Group & Base Infrastructure', level=2)
add_status_table([
    ("Create resource group RG_AutoFNOL_Logicapps (westus3)", "Succeeded", "Created for all new Logic App/App Service resources."),
    ("Resume Fabric capacity sbfabriccapeastus2 (East US 2)", "Succeeded", "Workspace WS_AutoFNOL's actual backing capacity; required resuming from paused state before Fabric SQL DB creation."),
])

add_heading('4.2 Fabric SQL Database (CaseThreadMap)', level=2)
add_status_table([
    ("Create Fabric SQL Database AutoFNOLCaseThreadMap in WS_AutoFNOL", "Succeeded", "Created via Fabric REST API (POST /v1/workspaces/{id}/sqlDatabases), polled LRO to Succeeded. Item id c08e3755-4fec-48d0-a173-035360cad306."),
    ("Apply CaseThreadMap schema (shared/case_thread_store_schema.sql)", "Succeeded", "Applied via pyodbc + Entra access-token auth (ODBC Driver 18). Verified table exists via INFORMATION_SCHEMA.TABLES query."),
])

add_heading('4.3 Work IQ Web Service (Graph tool bridge)', level=2)
add_status_table([
    ("Deploy as Azure Function (Consumption plan)", "Failed", "Blocked by a tenant-wide guardrail forcing publicNetworkAccess=Disabled on all new storage accounts, with no visible policy assignment to remove it; Azure Functions Consumption requires a storage account. Abandoned in favor of App Service."),
    ("Deploy as Azure App Service (Linux, Python 3.11, B1)", "Succeeded", "App app-autofnol-workiq created in RG_AutoFNOL_Logicapps; no storage account dependency, so unaffected by the guardrail above."),
    ("Create Entra app registration for Graph access (AutoFNOL-Logicapps-GraphApp)", "Succeeded", "App id c2b79763-6349-4080-ada9-14e80ee1fd1e; client secret generated (1yr); Graph app roles Sites.Read.All and Mail.Read granted (required a manual Graph API POST to /servicePrincipals/{id}/appRoleAssignments, since az ad app permission admin-consent reported success without actually creating the role assignments in this tenant)."),
    ("Test /healthz endpoint", "Succeeded", "Returns HTTP 200 {\"status\":\"ok\"}."),
    ("Test /search-documents endpoint (SharePoint/OneDrive via Graph Search)", "Succeeded", "Returns HTTP 200 with an empty result array (expected \u2013 no SharePoint test content has been configured yet); required adding a mandatory \"region\" field to the Graph Search API request body for app-only calls."),
    ("Test /search-mail endpoint (Outlook via Graph)", "Succeeded", "Returns HTTP 200 with real mail data from the monitored mailbox; required removing the $orderby parameter, which Graph does not allow combined with $search."),
])

add_heading('4.4 Fabric IQ Tool Integration (Ontology Data Agent)', level=2)
add_status_table([
    ("Build /query-ontology REST facade on Work IQ web app, proxying to the Fabric Data Agent MCP endpoint", "Succeeded", "Deployed; endpoint is reachable and returns HTTP 200 with a synthesized text answer."),
    ("Acquire a service-principal (app-only) Entra token for the Fabric MCP call", "Succeeded", "Token acquired successfully for both audiences tested (api.fabric.microsoft.com and analysis.windows.net/powerbi/api)."),
    ("Grant the service principal access to query the ontology data end-to-end", "Blocked", "The MCP transport call itself succeeds (isError=False) once the service principal was granted workspace Contributor on WS_AutoFNOL and added to the tenant's grp_spFabricAPIaccess security group. However, the Fabric Data Agent's own internal query against the ontology/Lakehouse still returns an authorization error to the agent, which the agent then paraphrases back as \u201cI'm not authorized to access the ontology graph.\u201d The same question succeeds immediately when tested with a delegated/interactive user token (via az login), confirming the data agent, ontology, and Lakehouse data themselves are all healthy \u2014 only the service-principal path is affected. Item-level sharing on the Ontology/Lakehouse items via the Fabric REST API returned \u201cFeatureNotAvailable \u2013 Item role sharing is not available,\u201d and direct SQL-endpoint queries against the same Lakehouse succeed fine for the same service principal, isolating the gap to the ontology/data-agent's own internal data-source binding. Root cause is most likely that the Fabric Data Agent's internal connection to its Lakehouse data source was configured using the creating user's own delegated OAuth connection at publish time, rather than a pass-through/organizational identity \u2014 which would require reconfiguring the data agent's data-source binding in the Fabric portal UI (not exposed via the REST APIs tried so far)."),
    ("Wire the Fabric IQ tool into the Foundry orchestrator agent as an OpenAPI tool", "Not yet run", "Code is ready (foundry/create_orchestrator_agent.py, foundry/fabric_ontology_openapi.json) but the orchestrator agent itself has not yet been created, pending resolution of the item above."),
])

add_heading('4.5 Foundry Orchestrator Agent', level=2)
add_status_table([
    ("Confirm SDK capabilities (azure-ai-agents vs azure-ai-projects)", "Succeeded", "azure-ai-agents==1.1.0 lacks a first-class MCP tool type; azure-ai-projects==2.5.0 has MCPTool/PromptAgentDefinition via the newer Responses API pattern. Chose to keep the simpler Assistants-style azure-ai-agents pattern with OpenAPI-tool facades for Fabric IQ and Work IQ."),
    ("Update Work IQ OpenAPI spec to point at the live App Service URL", "Succeeded", "foundry/workiq_graph_openapi.json updated to https://app-autofnol-workiq.azurewebsites.net."),
    ("Simplify orchestrator script's Work IQ tool auth to anonymous (no Foundry connection object required)", "Succeeded", "create_orchestrator_agent.py now uses OpenApiAnonymousAuthDetails instead of requiring a pre-created Foundry connection resource."),
    ("Create the orchestrator agent in Foundry project sbazureaimodels-project01", "Not yet run", "Pending resolution of the Fabric IQ service-principal authorization gap above, so that the orchestrator can be created with a working Fabric IQ tool from the outset."),
])

add_heading('4.6 Logic App (Trigger, Orchestration, Teams Posting)', level=2)
add_status_table([
    ("Deploy Logic App workflow (triggers/logicapp/workflow.json)", "Not yet run", "Not yet deployed; still references a plain Graph HTTP call for Teams posting that is known not to work (see next row)."),
    ("Validate Teams-posting approach", "Blocked", "Confirmed via az ad sp show that Microsoft Graph has no application-permission role for posting channel messages (only ChannelMessage.Read.All and ChannelMessage.UpdatePolicyViolation.All exist as Application-type roles; ChannelMessage.Send does not exist as an app-only permission at all). The current workflow.json's raw HTTP POST + Managed Identity auth to /teams/{id}/channels/{id}/messages will fail with the same \u201cBot hasn't been installed for user\u201d style error seen in the prior Copilot Studio-based accelerator. Needs to switch to the native Teams Logic App connector (delegated OAuth), matching the pattern that worked in the original solution."),
    ("Configure Office 365 Outlook trigger connection", "Not yet run", ""),
    ("Configure Teams connector connection (delegated OAuth)", "Not yet run", "Recommended remediation for the blocked item above."),
    ("Configure Fabric SQL connection for CaseThreadMap upsert", "Not yet run", ""),
])

add_heading('4.7 End-to-End Testing', level=2)
add_status_table([
    ("Send test email \u2192 confirm Logic App trigger fires", "Not yet run", ""),
    ("Confirm Foundry orchestrator agent runs and calls all 3 tools correctly", "Not yet run", ""),
    ("Confirm Teams channel message is posted with case triage summary", "Not yet run", ""),
    ("Confirm CaseThreadMap row is written correctly", "Not yet run", ""),
    ("Confirm a Teams reply/follow-up preserves the correct case context", "Not yet run", ""),
])

doc.add_page_break()

# ==================== 5. SUMMARY & NEXT STEPS ====================
add_heading('5. Summary & Recommended Next Steps', level=1)
doc.add_paragraph(
    "The data plane and state layer of this solution are fully deployed and verified: the Fabric SQL "
    "Database (CaseThreadMap) is live with its schema applied, and the Work IQ web service is live with "
    "all three of its tool endpoints (search-documents, search-mail, query-ontology) returning valid "
    "HTTP 200 responses. The single remaining blocker preventing the Foundry orchestrator agent from "
    "being created with full tool coverage is a Fabric Data Agent authorization gap that only affects "
    "service-principal (app-only) callers, not delegated/interactive users \u2014 despite the service "
    "principal being correctly granted workspace Contributor rights, Fabric API tenant access, and "
    "confirmed working access to the same Lakehouse data via its SQL endpoint."
)
doc.add_paragraph("Recommended next steps, in priority order:")
for i, s in enumerate([
    "Open the Fabric Data Agent (DA_AutoFNOL_Ontology) in the Fabric portal and inspect its configured "
    "data-source connection(s) under the agent's settings; if the data source shows a specific user's "
    "OAuth credential rather than \"organizational\"/pass-through auth, reconfigure or republish it to "
    "use an identity that supports service-principal callers.",
    "Once resolved, run foundry/create_orchestrator_agent.py to create the live orchestrator agent, "
    "capturing its agent ID for the Logic App to call.",
    "Rework the Logic App's Teams-posting action to use the native Teams connector with delegated OAuth "
    "(instead of the current raw Graph HTTP call), matching the working pattern from prior FNOL "
    "accelerator work.",
    "Deploy the Logic App with its three connections (Outlook trigger, Teams connector, Fabric SQL) and "
    "run the full end-to-end test: email \u2192 Foundry orchestrator \u2192 Teams post \u2192 CaseThreadMap write "
    "\u2192 Teams follow-up reply.",
    "Once end-to-end testing passes, produce the manual test-instructions document requested separately.",
]):
    doc.add_paragraph(f"{i+1}. {s}")

doc.save('FNOL_Logicapps_Solution_Status.docx')
print("Document generated: FNOL_Logicapps_Solution_Status.docx")
